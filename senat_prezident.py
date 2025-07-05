from pathlib import Path
from docx import Document
import chardet
import subprocess
import win32com.client
import re
import os

# === PHASE 1: Convert all .txt and .doc files to .docx ===

INPUT_DIR = Path("senat_zaznamy")
DOCX_DIR = Path("senat_zaznamy_docx")
DOCX_DIR.mkdir(exist_ok=True)

def convert_txt_to_docx(txt_path, docx_path):
    """
    Converts a plain text (.txt) file into a Word document (.docx),
    automatically detecting the character encoding of the input file.

    Args:
        txt_path (Path or str): Path to the input .txt file.
        docx_path (Path or str): Path where the output .docx should be saved.
    """
    try:
        with open(txt_path, "rb") as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)  # Detect the file's encoding (e.g., utf-8, windows-1250)
            encoding = result["encoding"]

        # Decode the text using detected encoding (fallback to utf-8)
        decoded_text = raw_data.decode(encoding or "utf-8", errors="replace")

        doc = Document()
        for line in decoded_text.splitlines():
            doc.add_paragraph(line)
        doc.save(docx_path)
        print(f"📄 TXT → DOCX: {txt_path.name} (encoding: {encoding})")
    except Exception as e:
        print(f"❌ Failed to convert TXT: {txt_path.name} → {e}")

def convert_doc_to_docx(doc_path, docx_path):
    """
    Converts a legacy Word (.doc) file to modern Word (.docx) format using LibreOffice.

    Requires LibreOffice (`soffice`) to be installed and available in system PATH.

    Args:
        doc_path (Path or str): Path to the input .doc file.
        docx_path (Path or str): Target path for the converted .docx file.
    """
    try:  # Call LibreOffice in headless mode to perform the conversion
        subprocess.run([
            "soffice", "--headless", "--convert-to", "docx",
            "--outdir", str(docx_path.parent),
            str(doc_path)
        ], check=True)
        print(f"Converted DOC: {doc_path.name} -> {docx_path.name}")
    except Exception as e:
        print(f"Failed DOC to DOCX: {doc_path.name} -> {e}")

def convert_all_to_docx(source_root="senat_zaznamy", target_root="senat_zaznamy_docx"):
    """
    Converts all .doc and .txt files under a source directory into .docx files,
    preserving the directory structure under a target directory.

    - .doc files are converted using Microsoft Word COM automation (Windows only)
    - .txt files are read and saved as simple .docx documents using python-docx

    Args:
        source_root (str): Root directory containing original .doc or .txt files.
        target_root (str): Target directory where .docx files will be saved.
    """
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    for root, _, files in os.walk(source_root):
        for file in files:
            full_path = os.path.join(root, file)
            rel_path = os.path.relpath(full_path, source_root)
            target_path = os.path.join(target_root, os.path.splitext(rel_path)[0] + ".docx")
            os.makedirs(os.path.dirname(target_path), exist_ok=True)

            if file.endswith(".doc") and not file.endswith(".docx"):
                try:
                    doc = word.Documents.Open(os.path.abspath(full_path))
                    doc.SaveAs(os.path.abspath(target_path), FileFormat=16)
                    doc.Close()
                    print(f"✅ DOC → DOCX: {rel_path}")
                except Exception as e:
                    print(f"❌ Failed to convert DOC: {rel_path} → {e}")

            elif file.endswith(".txt"):
                try:
                    with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()

                    docx = Document()
                    docx.add_paragraph(text)
                    docx.save(target_path)
                    print(f"📄 TXT → DOCX: {rel_path}")
                except Exception as e:
                    print(f"❌ Failed to convert TXT: {rel_path} → {e}")

    word.Quit()

# === PHASE 2: Analyze .docx files for "prezident" appearances ===
OUTPUT_FILE = DOCX_DIR / "senat_president.txt"

class SenatPresidentSearch:
    def __init__(self, input_dir="senat_zaznamy_docx"):
        self.input_dir = Path(input_dir)
        self.output_file = Path("senat_zaznamy") / "senat_prezident.txt"
        self.record_map = self.load_record_links()
        # print(self.record_map)
        self.results = []

    def load_record_links(self):
        record_path = Path("senat_zaznamy") / "records.txt"
        mapping = {}
        if not record_path.exists():
            print("⚠️ records.txt not found.")
            return mapping
        with open(record_path, "r", encoding="utf-8") as f:
            for line in f:
                if "\t" in line:
                    date, url = line.strip().split("\t", 1)
                    mapping[date] = url
        return mapping

    def extract_text_from_docx(self, file_path):
        """Extract text from .docx file"""
        try:
            doc = Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Error reading {file_path.name}: {e}")
            return ""

    def find_president_contexts(self, text):
        """Find all occurrences of 'prezident/president' with 2 words before and after"""
        # words = text.split()
        words = re.findall(r'\w+|[^\s\w]', text, flags=re.UNICODE)
        contexts = []
        
        for i, word in enumerate(words):
            clean_word = re.sub(r'[^\w]', '', word.lower())
            
            if 'prezident' in clean_word or 'president' in clean_word:
                if ":" in words[i+1:i+7] and 'č' in words[i+1]:
                    start = max(0, i - 2)
                    end = min(len(words), i + 3)
                    context = words[start:end]
                    contexts.append(' '.join(context))
        
        return contexts

    def search_files(self):
        """Search for 'prezident/president' in all .docx files"""
        print("Searching for 'prezident/president' in DOCX files...")
        
        for year_folder in self.input_dir.glob("*"):
            if year_folder.is_dir() and year_folder.name.isdigit():
                print(f"Searching in {year_folder.name}...")
                
                for file_path in year_folder.glob("*.docx"):
                    # Extract date from filename
                    date_match = re.search(r'(\d{4}-\d{2}-\d{2})', file_path.name)
                    if not date_match:
                        continue
                    
                    date = date_match.group(1)
                    text = self.extract_text_from_docx(file_path)
                    
                    if not text:
                        continue
                    
                    contexts = self.find_president_contexts(text)
                    
                    if contexts:
                        print(f"Found {len(contexts)} occurrences in {file_path.name}")
                        self.results.append({
                            'date': date,
                            'contexts': contexts
                        })

    def save_results(self):
        """Save results to senat_prezident.txt"""
        print(f"Saving results to {self.output_file}...")
        
        self.results.sort(key=lambda x: x['date'])
        
        with open(self.output_file, 'w', encoding='utf-8') as f:
            for result in self.results:
                contexts_str = ' | '.join(result['contexts'])
                link = self.record_map.get(result['date'], "NO_LINK")
                f.write(f"{result['date']}\t{link}\t{contexts_str}\n")
        
        print(f"Saved {len(self.results)} files with prezident mentions")

    def run(self):
        """Main search function"""
        self.search_files()
        if self.results:
            self.save_results()
            print(f"Complete: Found prezident in {len(self.results)} files")
        else:
            print("No occurrences of 'prezident/president' found")

# === RUN BOTH PHASES ===
if __name__ == "__main__":
    # convert_all_to_docx()
    searcher = SenatPresidentSearch()
    searcher.run()

